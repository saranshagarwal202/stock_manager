from numpy import nan
from MainWindow import Ui_MainWindow

from PyQt5 import QtWidgets as qtw
from PyQt5 import QtCore as qtc
from PyQt5 import QtGui as qtg
from PyQt5 import QtPrintSupport as qtp

from copy import copy, deepcopy
import pandas as pd
import json
import os
from datetime import date
import docx


class MainWindow(qtw.QMainWindow):

    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)

        self.ui = Ui_MainWindow()
        self.ui.setupUi(self)
        self.ui.actionSave_As.setEnabled(False)
        self.ui.actionSave.setEnabled(False)
        self.ui.tabWidget.setTabEnabled(3, False)
        self.ui.tabWidget.setTabEnabled(2, False)
        self.ui.tabWidget.setTabEnabled(1, False)
        self.ui.tabWidget.setTabEnabled(0, False)

        # Push buttons
        ##### Stock Sheet #####
        # self.ui.edit_button.setEnabled(False)
        self.ui.clear_button.setEnabled(False)
        self.ui.delete_entry.setEnabled(False)
        self.ui.reset_find_preview.setEnabled(False)
        self.ui.confirm_edit.setEnabled(False)
        self.ui.confirm_edit.clicked.connect(
            lambda: self.confirm_edit(tab="stock"))
        self.ui.Add_stock.clicked.connect(self.add_stock)
        self.ui.find_button.clicked.connect(
            lambda: self.find_item_in_stock_sheet(sheet="stock_sheet"))
        self.ui.delete_entry.clicked.connect(
            lambda: self.delete_entry_from_tabel(tab="stock"))
        self.ui.clear_button.clicked.connect(
            lambda: self.clear_tabel(self.ui.current_stock_find_entry))
        self.ui.clear_button.clicked.connect(
            lambda: self.clear_tabel(self.ui.find_stock_preview, True))
        self.ui.clear_button.clicked.connect(self.stock_tab_button_set_False)
        self.ui.reset_find_preview.clicked.connect(
            lambda: self.reset_preview(tab="stock"))
        self.ui.current_stock_find_entry.keyPressEvent = self.find_enter_event
        self.ui.find_stock_preview.keyPressEvent = self.preview_enter_event
        self.ui.scrollArea_2.mousePressEvent = self.reset_selected

        ##### Issue Sheet #####
        self.ui.clear_button_2.setEnabled(False)
        self.ui.delete_entry_2.setEnabled(False)
        self.ui.edit_button_2.setEnabled(False)
        self.ui.Reset_issued_preview.setEnabled(False)
        self.ui.approval_issue.setEnabled(False)
        self.ui.reset_issue_items.setEnabled(False)
        self.ui.Delete_issue_items.setEnabled(False)
        self.ui.add_sku.setEnabled(False)
        self.ui.Weight_to_add.setEnabled(False)
        self.ui.find_button_2.clicked.connect(
            lambda: self.find_item_in_stock_sheet(sheet="issue_sheet"))
        self.ui.clear_button_2.clicked.connect(
            lambda: self.clear_tabel(self.ui.issued_stock_find_entry))
        self.ui.clear_button_2.clicked.connect(
            lambda: self.clear_tabel(self.ui.issue_stock_preview, True))
        self.ui.clear_button_2.clicked.connect(self.issue_tab_button_set_false)
        self.ui.edit_button_2.clicked.connect(
            lambda: self.confirm_edit(tab="issue"))
        self.ui.Reset_issued_preview.clicked.connect(
            lambda: self.reset_preview(tab="issue"))
        self.ui.delete_entry_2.clicked.connect(
            lambda: self.delete_entry_from_tabel(tab="issue"))
        self.ui.find_sku.clicked.connect(self.find_sku)
        self.ui.clear_issue_stock.clicked.connect(
            lambda: self.clear_tabel(self.ui.sku_find, True))
        self.ui.clear_issue_stock.clicked.connect(
            lambda: self.ui.add_sku.setEnabled(False))
        self.ui.clear_issue_stock.clicked.connect(
            lambda: self.ui.enter_sku_code.clear())
        self.ui.clear_issue_stock.clicked.connect(
            lambda: self.ui.Weight_to_add.clear())
        self.ui.clear_issue_stock.clicked.connect(
            lambda: self.ui.Weight_to_add.setMaximum(0))
        self.ui.Delete_issue_items.clicked.connect(
            self.delete_item_from_issue_table)
        self.ui.reset_issue_items.clicked.connect(self.reset_issue_items)
        self.ui.approval_issue.clicked.connect(self.issue_approval)
        # self.ui.clear_issue_stock.clicked.connect(
        #     lambda: self.ui.Weight_to_add.clear())
        self.ui.add_sku.clicked.connect(self.add_sku)
        self.ui.dateEdit_2.setDate(
            qtc.QDate(date.today().year, date.today().month, date.today().day))
        self.ui.scrollArea.mousePressEvent = self.reset_selected_issue

        # Issue return sheet
        self.ui.show_issue_return.clicked.connect(self.show_issue_return)
        self.ui.dateEdit.setDate(
            qtc.QDate(date.today().year, date.today().month, date.today().day))
        self.ui.issue_returned.clicked.connect(self.issue_returned)
        self.ui.issue_returned.setEnabled(False)
        self.ui.reset_issue_return.setEnabled(False)
        self.ui.reset_issue_return.clicked.connect(self.reset_issue_return)

        # Sale tab
        self.ui.clear_sale_table.setEnabled(False)
        self.ui.sale_button.setEnabled(False)
        self.ui.calculate_sale.setEnabled(False)
        self.ui.show_sale_2.clicked.connect(self.show_sale)
        self.ui.dateEdit_3.setDate(
            qtc.QDate(date.today().year, date.today().month, date.today().day))
        self.ui.sale_button.clicked.connect(self.make_sale)
        self.ui.clear_sale_table.clicked.connect(
            lambda: self.ui.broker_name_sale.clear())
        self.ui.clear_sale_table.clicked.connect(
            lambda: self.ui.party_name_sale.clear())
        self.ui.clear_sale_table.clicked.connect(
            lambda: self.clear_tabel(self.ui.sale_table, row_count=0))
        self.ui.calculate_sale.clicked.connect(self.calculate_sale)

        # tabel
        # self.ui.find_stock_preview.clicked.connect(lambda : self.ui.delete_entry.setEnabled(True))

        # Variabels
        self.saved = True
        self.deleted = []
        self.original = []
        self.original_issue = []
        self.sheet_index_add_sku = None
        self.issue_sheet_skus = []
        # self.original_prev = []

        # Action buttons
        self.ui.actionOpen_Project.triggered.connect(self.open_project)
        self.ui.actionExit_2.triggered.connect(self.exit)
        self.ui.actionNew_Project.triggered.connect(self.New_project)
        self.ui.actionSave.triggered.connect(
            lambda: self.save_project(save_path=self.project_path))
        self.ui.actionSave_As.triggered.connect(lambda: self.save_project(qtw.QFileDialog.getExistingDirectory(
            self, "Select Directory")))

    def red_palette(self):
        palette = qtg.QPalette()
        brush = qtg.QBrush(qtg.QColor(0, 0, 0))
        brush.setStyle(qtc.Qt.SolidPattern)
        palette.setBrush(qtg.QPalette.Active, qtg.QPalette.WindowText, brush)
        brush = qtg.QBrush(qtg.QColor(239, 41, 41))
        brush.setStyle(qtc.Qt.SolidPattern)
        palette.setBrush(qtg.QPalette.Active, qtg.QPalette.Button, brush)
        brush = qtg.QBrush(qtg.QColor(255, 147, 147))
        brush.setStyle(qtc.Qt.SolidPattern)
        palette.setBrush(qtg.QPalette.Active, qtg.QPalette.Light, brush)
        brush = qtg.QBrush(qtg.QColor(247, 94, 94))
        brush.setStyle(qtc.Qt.SolidPattern)
        palette.setBrush(qtg.QPalette.Active, qtg.QPalette.Midlight, brush)
        brush = qtg.QBrush(qtg.QColor(119, 20, 20))
        brush.setStyle(qtc.Qt.SolidPattern)
        palette.setBrush(qtg.QPalette.Active, qtg.QPalette.Dark, brush)
        brush = qtg.QBrush(qtg.QColor(159, 27, 27))
        brush.setStyle(qtc.Qt.SolidPattern)
        palette.setBrush(qtg.QPalette.Active, qtg.QPalette.Mid, brush)
        brush = qtg.QBrush(qtg.QColor(0, 0, 0))
        brush.setStyle(qtc.Qt.SolidPattern)
        palette.setBrush(qtg.QPalette.Active, qtg.QPalette.Text, brush)
        brush = qtg.QBrush(qtg.QColor(255, 255, 255))
        brush.setStyle(qtc.Qt.SolidPattern)
        palette.setBrush(qtg.QPalette.Active, qtg.QPalette.BrightText, brush)
        brush = qtg.QBrush(qtg.QColor(0, 0, 0))
        brush.setStyle(qtc.Qt.SolidPattern)
        palette.setBrush(qtg.QPalette.Active, qtg.QPalette.ButtonText, brush)
        brush = qtg.QBrush(qtg.QColor(255, 255, 255))
        brush.setStyle(qtc.Qt.SolidPattern)
        palette.setBrush(qtg.QPalette.Active, qtg.QPalette.Base, brush)
        brush = qtg.QBrush(qtg.QColor(239, 41, 41))
        brush.setStyle(qtc.Qt.SolidPattern)
        palette.setBrush(qtg.QPalette.Active, qtg.QPalette.Window, brush)
        brush = qtg.QBrush(qtg.QColor(0, 0, 0))
        brush.setStyle(qtc.Qt.SolidPattern)
        palette.setBrush(qtg.QPalette.Active, qtg.QPalette.Shadow, brush)
        brush = qtg.QBrush(qtg.QColor(247, 148, 148))
        brush.setStyle(qtc.Qt.SolidPattern)
        palette.setBrush(qtg.QPalette.Active,
                         qtg.QPalette.AlternateBase, brush)
        brush = qtg.QBrush(qtg.QColor(255, 255, 220))
        brush.setStyle(qtc.Qt.SolidPattern)
        palette.setBrush(qtg.QPalette.Active, qtg.QPalette.ToolTipBase, brush)
        brush = qtg.QBrush(qtg.QColor(0, 0, 0))
        brush.setStyle(qtc.Qt.SolidPattern)
        palette.setBrush(qtg.QPalette.Active, qtg.QPalette.ToolTipText, brush)
        brush = qtg.QBrush(qtg.QColor(0, 0, 0, 128))
        brush.setStyle(qtc.Qt.SolidPattern)
        palette.setBrush(qtg.QPalette.Active,
                         qtg.QPalette.PlaceholderText, brush)
        brush = qtg.QBrush(qtg.QColor(0, 0, 0))
        brush.setStyle(qtc.Qt.SolidPattern)
        palette.setBrush(qtg.QPalette.Inactive, qtg.QPalette.WindowText, brush)
        brush = qtg.QBrush(qtg.QColor(239, 41, 41))
        brush.setStyle(qtc.Qt.SolidPattern)
        palette.setBrush(qtg.QPalette.Inactive, qtg.QPalette.Button, brush)
        brush = qtg.QBrush(qtg.QColor(255, 147, 147))
        brush.setStyle(qtc.Qt.SolidPattern)
        palette.setBrush(qtg.QPalette.Inactive, qtg.QPalette.Light, brush)
        brush = qtg.QBrush(qtg.QColor(247, 94, 94))
        brush.setStyle(qtc.Qt.SolidPattern)
        palette.setBrush(qtg.QPalette.Inactive, qtg.QPalette.Midlight, brush)
        brush = qtg.QBrush(qtg.QColor(119, 20, 20))
        brush.setStyle(qtc.Qt.SolidPattern)
        palette.setBrush(qtg.QPalette.Inactive, qtg.QPalette.Dark, brush)
        brush = qtg.QBrush(qtg.QColor(159, 27, 27))
        brush.setStyle(qtc.Qt.SolidPattern)
        palette.setBrush(qtg.QPalette.Inactive, qtg.QPalette.Mid, brush)
        brush = qtg.QBrush(qtg.QColor(0, 0, 0))
        brush.setStyle(qtc.Qt.SolidPattern)
        palette.setBrush(qtg.QPalette.Inactive, qtg.QPalette.Text, brush)
        brush = qtg.QBrush(qtg.QColor(255, 255, 255))
        brush.setStyle(qtc.Qt.SolidPattern)
        palette.setBrush(qtg.QPalette.Inactive, qtg.QPalette.BrightText, brush)
        brush = qtg.QBrush(qtg.QColor(0, 0, 0))
        brush.setStyle(qtc.Qt.SolidPattern)
        palette.setBrush(qtg.QPalette.Inactive, qtg.QPalette.ButtonText, brush)
        brush = qtg.QBrush(qtg.QColor(255, 255, 255))
        brush.setStyle(qtc.Qt.SolidPattern)
        palette.setBrush(qtg.QPalette.Inactive, qtg.QPalette.Base, brush)
        brush = qtg.QBrush(qtg.QColor(239, 41, 41))
        brush.setStyle(qtc.Qt.SolidPattern)
        palette.setBrush(qtg.QPalette.Inactive, qtg.QPalette.Window, brush)
        brush = qtg.QBrush(qtg.QColor(0, 0, 0))
        brush.setStyle(qtc.Qt.SolidPattern)
        palette.setBrush(qtg.QPalette.Inactive, qtg.QPalette.Shadow, brush)
        brush = qtg.QBrush(qtg.QColor(247, 148, 148))
        brush.setStyle(qtc.Qt.SolidPattern)
        palette.setBrush(qtg.QPalette.Inactive,
                         qtg.QPalette.AlternateBase, brush)
        brush = qtg.QBrush(qtg.QColor(255, 255, 220))
        brush.setStyle(qtc.Qt.SolidPattern)
        palette.setBrush(qtg.QPalette.Inactive,
                         qtg.QPalette.ToolTipBase, brush)
        brush = qtg.QBrush(qtg.QColor(0, 0, 0))
        brush.setStyle(qtc.Qt.SolidPattern)
        palette.setBrush(qtg.QPalette.Inactive,
                         qtg.QPalette.ToolTipText, brush)
        brush = qtg.QBrush(qtg.QColor(0, 0, 0, 128))
        brush.setStyle(qtc.Qt.SolidPattern)
        palette.setBrush(qtg.QPalette.Inactive,
                         qtg.QPalette.PlaceholderText, brush)
        brush = qtg.QBrush(qtg.QColor(119, 20, 20))
        brush.setStyle(qtc.Qt.SolidPattern)
        palette.setBrush(qtg.QPalette.Disabled, qtg.QPalette.WindowText, brush)
        brush = qtg.QBrush(qtg.QColor(239, 41, 41))
        brush.setStyle(qtc.Qt.SolidPattern)
        palette.setBrush(qtg.QPalette.Disabled, qtg.QPalette.Button, brush)
        brush = qtg.QBrush(qtg.QColor(255, 147, 147))
        brush.setStyle(qtc.Qt.SolidPattern)
        palette.setBrush(qtg.QPalette.Disabled, qtg.QPalette.Light, brush)
        brush = qtg.QBrush(qtg.QColor(247, 94, 94))
        brush.setStyle(qtc.Qt.SolidPattern)
        palette.setBrush(qtg.QPalette.Disabled, qtg.QPalette.Midlight, brush)
        brush = qtg.QBrush(qtg.QColor(119, 20, 20))
        brush.setStyle(qtc.Qt.SolidPattern)
        palette.setBrush(qtg.QPalette.Disabled, qtg.QPalette.Dark, brush)
        brush = qtg.QBrush(qtg.QColor(159, 27, 27))
        brush.setStyle(qtc.Qt.SolidPattern)
        palette.setBrush(qtg.QPalette.Disabled, qtg.QPalette.Mid, brush)
        brush = qtg.QBrush(qtg.QColor(119, 20, 20))
        brush.setStyle(qtc.Qt.SolidPattern)
        palette.setBrush(qtg.QPalette.Disabled, qtg.QPalette.Text, brush)
        brush = qtg.QBrush(qtg.QColor(255, 255, 255))
        brush.setStyle(qtc.Qt.SolidPattern)
        palette.setBrush(qtg.QPalette.Disabled, qtg.QPalette.BrightText, brush)
        brush = qtg.QBrush(qtg.QColor(119, 20, 20))
        brush.setStyle(qtc.Qt.SolidPattern)
        palette.setBrush(qtg.QPalette.Disabled, qtg.QPalette.ButtonText, brush)
        brush = qtg.QBrush(qtg.QColor(239, 41, 41))
        brush.setStyle(qtc.Qt.SolidPattern)
        palette.setBrush(qtg.QPalette.Disabled, qtg.QPalette.Base, brush)
        brush = qtg.QBrush(qtg.QColor(239, 41, 41))
        brush.setStyle(qtc.Qt.SolidPattern)
        palette.setBrush(qtg.QPalette.Disabled, qtg.QPalette.Window, brush)
        brush = qtg.QBrush(qtg.QColor(0, 0, 0))
        brush.setStyle(qtc.Qt.SolidPattern)
        palette.setBrush(qtg.QPalette.Disabled, qtg.QPalette.Shadow, brush)
        brush = qtg.QBrush(qtg.QColor(239, 41, 41))
        brush.setStyle(qtc.Qt.SolidPattern)
        palette.setBrush(qtg.QPalette.Disabled,
                         qtg.QPalette.AlternateBase, brush)
        brush = qtg.QBrush(qtg.QColor(255, 255, 220))
        brush.setStyle(qtc.Qt.SolidPattern)
        palette.setBrush(qtg.QPalette.Disabled,
                         qtg.QPalette.ToolTipBase, brush)
        brush = qtg.QBrush(qtg.QColor(0, 0, 0))
        brush.setStyle(qtc.Qt.SolidPattern)
        palette.setBrush(qtg.QPalette.Disabled,
                         qtg.QPalette.ToolTipText, brush)
        brush = qtg.QBrush(qtg.QColor(0, 0, 0, 128))
        brush.setStyle(qtc.Qt.SolidPattern)
        palette.setBrush(qtg.QPalette.Disabled,
                         qtg.QPalette.PlaceholderText, brush)

        return palette

    def show_sale(self):
        approval_no = self.ui.approval_no_3.value()
        self.sale_approval_list = self.issue_sheet.loc[self.issue_sheet["APPROVAL NO."] == approval_no]
        self.sale_approval_list.drop("ASKING RATE", axis=1, inplace=True)

        already_returned_list = self.issue_return_sheet.loc[
            self.issue_return_sheet["APPROVAL NO."] == approval_no]
        already_returned_list = already_returned_list.groupby("SKU", as_index=False)[
            "Weight returned"].sum()
        already_sold_list = self.sale_sheet.loc[self.sale_sheet["APPROVAL NO."] == approval_no]
        already_sold_list = already_sold_list.groupby(
            "SKU", as_index=False)["Qty Sold"].sum()

        def apply_func(series):
            already_returned_weight = 0
            if series["SKU"] in already_returned_list["SKU"].values:
                already_returned_weight = already_returned_list.loc[
                    already_returned_list["SKU"] == series["SKU"]].iloc[0]["Weight returned"]
            if series["SKU"] in already_sold_list["SKU"].values:
                already_returned_weight += already_sold_list.loc[already_sold_list["SKU"]
                                                                 == series["SKU"]].iloc[0]["Qty Sold"]
            series["Sale Date"] = ''
            qspinbox = qtw.QDoubleSpinBox()
            qspinbox.setMaximum(
                series["Weight issued"]-already_returned_weight)
            qspinbox.setMinimum(0.01)
            qspinbox.setDecimals(2)
            qspinbox.setValue(
                series["Weight issued"] - already_returned_weight)
            qspinbox.setToolTip(
                f'total weight available = weight issued to this approval{series["Weight issued"]} - already returned/sold weight ({already_returned_weight})')
            series["Qty Sold"] = qspinbox

            qspinbox2 = qtw.QDoubleSpinBox()
            qspinbox2.setDecimals(2)
            qspinbox2.setMaximum(99999999)
            series["Selling Price"] = qspinbox2
            series["Amount"] = qspinbox2.value()*qspinbox.value()

            qspinbox3 = qtw.QDoubleSpinBox()
            qspinbox3.setDecimals(2)
            qspinbox3.setMaximum(100)
            series["Discount (%)"] = qspinbox3
            series["Net amount after discount"] = series["Amount"] - \
                (series["Amount"]*qspinbox3.value()/100)

            qspinbox4 = qtw.QDoubleSpinBox()
            qspinbox4.setDecimals(2)
            qspinbox4.setMaximum(100)
            series["Brokerage (%)"] = qspinbox4
            series["Brokerage amount"] = series["Net amount after discount"] * \
                qspinbox4.value()/100

            series["Net amount after discount and brokerage"] = series["Net amount after discount"] - \
                series["Brokerage amount"]

            series["Payment Status"] = ''

            qcombobox = qtw.QComboBox()
            qcombobox.addItems(['No', 'Yes'])
            qcombobox.setCurrentText(series["SELECTION YES/NO"])
            series["SELECTION YES/NO"] = qcombobox
            return series

        self.sale_approval_list = self.sale_approval_list.apply(
            lambda x: apply_func(x), axis=1)
        self.sale_approval_list = self.sale_approval_list[self.schema["sale_sheet"]["columns"]]

        if not self.sale_approval_list.empty:
            self.load_data_in_tabel(
                self.ui.sale_table, self.sale_approval_list,
                len(
                    self.sale_approval_list), self.schema["sale_sheet"]["columns"],
                read_only=[
                    "Approval Date", "Sale Date", "APPROVAL NO.", "SKU", "PL No.",
                    "Item", "Shape", "Size", "Mine", "Grade", "Price Code", "Lot Details (If Any)",
                    "Original weight", "Weight issued",
                    "Amount", "Net amount after discount", "Brokerage amount",
                    "Net amount after discount and brokerage", "BROKER NAME", "PARTY NAME"],
                set_cell_widget=["Qty Sold", "Selling Price", "Discount (%)", "Brokerage (%)", "SELECTION YES/NO"])
            self.ui.clear_sale_table.setEnabled(True)
            self.ui.sale_button.setEnabled(True)
            self.ui.calculate_sale.setEnabled(True)
            # set broker name and party name as per first item in approval
            try:
                self.ui.broker_name_sale.setText(
                    self.sale_approval_list.iloc[0]["BROKER NAME"])
            except:
                # Unable to set because none
                False
            try:
                self.ui.party_name_sale.setText(
                    self.sale_approval_list.iloc[0]["PARTY NAME"])
            except:
                False

    def make_sale(self):
        # check if broker name and party name is filled
        passed = True
        if bool(self.ui.broker_name_sale.text()):
            self.ui.broker_name_sale.setPalette(self.style().standardPalette())

        else:
            self.ui.broker_name_sale.setPalette(self.red_palette())
            passed = False

        if bool(self.ui.party_name_sale.text()):
            self.ui.party_name_sale.setPalette(self.style().standardPalette())
        else:
            self.ui.party_name_sale.setPalette(self.red_palette())
            passed = False

        if passed:
            sale_date = self.ui.dateEdit_3.text()
            broker_name = self.ui.broker_name_sale.text()
            party_name = self.ui.party_name_sale.text()
            # calculate amount before sale
            self.calculate_sale()
            approval_no = self.ui.approval_no_3.value()
            already_returned_list = self.issue_return_sheet.loc[
                self.issue_return_sheet["APPROVAL NO."] == approval_no]
            already_returned_list = already_returned_list.groupby("SKU", as_index=False)[
                "Weight returned"].sum()
            already_sold_list = self.sale_sheet.loc[self.sale_sheet["APPROVAL NO."] == approval_no]
            already_sold_list = already_sold_list.groupby(
                "SKU", as_index=False)["Qty Sold"].sum()

            # for dicti in self.sale_approval_list:
            def apply_func(series):
                already_returned_weight = 0
                if series["SKU"] in already_returned_list["SKU"].values:
                    already_returned_weight = already_returned_list.loc[
                        already_returned_list["SKU"] == series["SKU"]].iloc[0]["Weight returned"]
                if series["SKU"] in already_sold_list["SKU"].values:
                    already_returned_weight += already_sold_list.loc[already_sold_list["SKU"]
                                                                     == series["SKU"]].iloc[0]["Qty Sold"]
                series["BROKER NAME"] = broker_name
                series["PARTY NAME"] = party_name
                series["Sale Date"] = sale_date
                series["SELECTION YES/NO"] = series["SELECTION YES/NO"].currentText()
                series["Qty Sold"] = series["Qty Sold"].value()
                series["Selling Price"] = series["Selling Price"].value()
                series["Discount (%)"] = series["Discount (%)"].value()
                series["Brokerage (%)"] = series["Brokerage (%)"].value()

                self.stock_sheet.loc[self.stock_sheet["SKU"] == series["SKU"], [
                    "Sold Qty", "Weight in Hand"]] += [series["Qty Sold"], -series["Qty Sold"]]
                self.issue_sheet.loc[self.issue_sheet["SKU"] == series["SKU"], [
                    "Qty Sold"]] = [series["Qty Sold"]]
                self.issue_sheet.drop(
                    self.issue_sheet[self.issue_sheet["Qty Sold"]+already_returned_weight >= series["Weight issued"]].index, inplace=True)
                return series

            self.sale_approval_list = self.sale_approval_list.apply(
                lambda x: apply_func(x), axis=1)
            self.sale_sheet = self.sale_sheet.append(
                self.sale_approval_list, ignore_index=True)
            # self.sale_sheet.extend(self.sale_approval_list)
            # Show pop up that sale is finished
            msg = qtw.QMessageBox()
            msg.setText(
                f"Sale of {len(self.sale_approval_list)} item(s) added to sale sheet in the name of broker {broker_name} and party {party_name} on {sale_date}")
            msg.setIcon(qtw.QMessageBox.Information)
            msg.exec_()
            self.ui.clear_sale_table.click()
            total_sale = {
                "Qty Sold": 0,
                "Amount": 0,
                "Net amount after discount": 0,
                "Brokerage amount": 0,
                "Net amount after discount and brokerage": 0
            }
            for i, key in enumerate(total_sale):
                self.ui.total_sale.setItem(i, 0, qtw.QTableWidgetItem(
                    str(total_sale[key])))
            self.ui.total_sale.horizontalHeader(
            ).resizeSections(qtw.QHeaderView.ResizeToContents)

    def calculate_sale(self):
        sale_date = self.ui.dateEdit_3.text()
        broker_name = self.ui.broker_name_sale.text()
        party_name = self.ui.party_name_sale.text()
        total_sale = {
            "Qty Sold": 0,
            "Amount": 0,
            "Net amount after discount": 0,
            "Brokerage amount": 0,
            "Net amount after discount and brokerage": 0
        }

        def apply_func(series):
            series["Sale Date"] = sale_date
            series["BROKER NAME"] = broker_name
            series["PARTY NAME"] = party_name

            series["Amount"] = series["Qty Sold"].value() * \
                series["Selling Price"].value()
            series["Net amount after discount"] = series["Amount"] - \
                (series["Amount"]*(series["Discount (%)"].value()/100))
            series["Brokerage amount"] = series["Net amount after discount"] * \
                (series["Brokerage (%)"].value()/100)
            series["Net amount after discount and brokerage"] = series["Net amount after discount"] - \
                series["Brokerage amount"]

            total_sale["Qty Sold"] += series["Qty Sold"].value()
            total_sale["Amount"] += series["Amount"]
            total_sale["Brokerage amount"] += series["Brokerage amount"]
            total_sale["Net amount after discount"] += series["Net amount after discount"]
            total_sale["Net amount after discount and brokerage"] += series["Net amount after discount and brokerage"]
            return series
        self.sale_approval_list = self.sale_approval_list.apply(
            lambda x: apply_func(x), axis=1)
        # load data in total sale
        self.ui.total_sale.setRowCount(len(total_sale))
        self.ui.total_sale.setColumnCount(1)
        self.ui.total_sale.setHorizontalHeaderLabels(['Total'])
        self.ui.total_sale.setVerticalHeaderLabels(list(total_sale))
        for i, key in enumerate(total_sale):
            self.ui.total_sale.setItem(i, 0, qtw.QTableWidgetItem(
                str(total_sale[key])))
        self.ui.total_sale.horizontalHeader(
        ).resizeSections(qtw.QHeaderView.ResizeToContents)

        self.load_data_in_tabel(
            self.ui.sale_table, self.sale_approval_list,
            len(
                self.sale_approval_list), self.schema["sale_sheet"]["columns"],
            read_only=[
                "Approval Date", "Sale Date", "APPROVAL NO.", "SKU", "PL No.",
                "Item", "Shape", "Size", "Mine", "Grade", "Price Code", "Lot Details (If Any)",
                "Original weight", "Weight issued",
                "Amount", "Net amount after discount", "Brokerage amount",
                "Net amount after discount and brokerage", "BROKER NAME", "PARTY NAME"],
            set_cell_widget=["Qty Sold", "Selling Price", "Discount (%)", "Brokerage (%)", "SELECTION YES/NO"])

    def issue_returned(self):
        # found_list = []
        # sheet_index = 0

        def apply_func(series):
            series["Short/Excess Weight"] = series["Short/Excess Weight"].value()

            if series["Short/Excess Weight"] < 0:
                issue_qty = -(series["Weight returned"].value() -
                              series["Short/Excess Weight"])
                weight_in_hand = series["Weight returned"].value()
                series["complete_return"] = (series["Weight returned"].value(
                ) - series["Short/Excess Weight"]) >= series["Weight returned"].maximum()
            else:
                issue_qty = -series["Weight returned"].value()
                weight_in_hand = series["Weight returned"].value() + \
                    series["Short/Excess Weight"]
                series["complete_return"] = series["Weight returned"].value(
                ) >= series["Weight returned"].maximum()
            series["Weight returned"] = series["Weight returned"].value()
            self.stock_sheet.loc[self.stock_sheet["SKU"] == series["SKU"], ["Weight in Hand", "Short/Excess Received", "Issue Qty"]] += [
                weight_in_hand, series["Short/Excess Weight"], issue_qty]
            self.load_data_in_tabel(
                self.ui.current_stock, values=self.stock_sheet,
                rowcount=self.ui.current_stock.rowCount(),
                columns=self.schema["stock_sheet"]["columns"],
                affected_row=self.stock_sheet[self.stock_sheet["SKU"] == series["SKU"]].index.tolist()[0])
            return series

        self.issue_approval_list = self.issue_approval_list.apply(
            lambda x: apply_func(x), axis=1)
        complete_return_items = self.issue_approval_list.loc[
            self.issue_approval_list["complete_return"]]
        if not complete_return_items.empty:
            complete_return_items.drop("complete_return", axis=1, inplace=True)
            # self.issue_approval_list.drop(
            #     complete_return_items.index, inplace=True)
            # self.issue_return_sheet = self.issue_return_sheet.append(
            #     complete_return_items, ignore_index=True)
            self.issue_sheet.drop(
                self.issue_sheet.loc[self.issue_sheet["SKU"].isin(complete_return_items["SKU"])].index, inplace=True)

        self.issue_approval_list.drop("complete_return", axis=1, inplace=True)
        self.issue_approval_list.drop(
            self.issue_approval_list[self.issue_approval_list["Weight returned"] == 0].index, inplace=True)
        self.issue_return_sheet = self.issue_return_sheet.append(
            self.issue_approval_list, ignore_index=True)

        # self.issue_sheet.loc[self.issue_sheet["SKU"].isin(
        #                     self.issue_approval_list["SKU"])] = self.issue_approval_list
        # for dicti in self.stock_sheet:
        #     count = 0
        #     for row in self.issue_approval_list:
        #         row["Weight returned"] = row["Weight returned"].value()

        #         if dicti["SKU"] == row["SKU"]:
        #             # check if this approval also exists in issue_return_sheet
        #             already_returned_weight = 0
        #             i = 0
        #             for return_dicti in self.issue_return_sheet:
        #                 if return_dicti["APPROVAL NO."] == row["APPROVAL NO."]:
        #                     if return_dicti["SKU"] == row["SKU"]:
        #                         already_returned_weight = return_dicti["Weight returned"]
        #                         del(self.issue_return_sheet[i])
        #                         break
        #                 i += 1

        #             dicti["Weight in Hand"] += row["Weight returned"]
        #             dicti["Issue Qty"] -= row["Weight returned"]
        #             self.load_data_in_tabel(
        #                 self.ui.current_stock, values=dicti,
        #                 rowcount=self.ui.current_stock.rowCount(),
        #                 columns=self.schema["stock_sheet"]["columns"],
        #                 affected_row=sheet_index)
        #             row["Weight returned"] += already_returned_weight
        #             # delete item from issue_sheet if all weight is returned
        #             if row["Weight returned"] == 0:
        #                 issue_i = 0
        #                 for issue_dicti in self.issue_sheet:
        #                     if issue_dicti["APPROVAL NO."] == row["APPROVAL NO."]:
        #                         if issue_dicti["SKU"] == issue_dicti["SKU"]:
        #                             break
        #                     issue_i += 1
        #                 del(self.issue_sheet[issue_i])

        #             found_list.append(self.issue_approval_list.pop(count))
        #             break
        #         count += 1
        #     sheet_index += 1
        # assert (bool(self.issue_approval_list) == False)
        # self.issue_return_sheet.extend(found_list)
        self.reset_issue_return()

    def reset_issue_return(self):
        self.ui.issue_returned.setEnabled(False)
        self.ui.reset_issue_return.setEnabled(False)
        self.ui.issue_return_table.clear()
        self.ui.issue_return_table.setRowCount(0)
        self.ui.issue_return_table.setColumnCount(
            len(self.schema["issue_return_sheet"]["columns"]))
        self.ui.issue_return_table.setHorizontalHeaderLabels(
            self.schema["issue_return_sheet"]["columns"])

    def show_issue_return(self):
        approval_no = self.ui.approval_no.value()
        return_date = self.ui.dateEdit.text()

        self.issue_approval_list = self.issue_sheet.loc[self.issue_sheet["APPROVAL NO."] == approval_no]
        if self.issue_approval_list.empty:
            return
        already_returned_list = self.issue_return_sheet.loc[
            self.issue_return_sheet["APPROVAL NO."] == approval_no]
        already_returned_list = already_returned_list.groupby("SKU", as_index=False)[
            "Weight returned"].sum()
        already_sold_list = self.sale_sheet.loc[self.sale_sheet["APPROVAL NO."] == approval_no]
        already_sold_list = already_sold_list.groupby(
            "SKU", as_index=False)["Qty Sold"].sum()

        def apply_func(series):
            already_returned_weight = 0
            if series["SKU"] in already_returned_list["SKU"].values:
                already_returned_weight = already_returned_list.loc[
                    already_returned_list["SKU"] == series["SKU"]].iloc[0]["Weight returned"]
            if series["SKU"] in already_sold_list["SKU"].values:
                already_returned_weight += already_sold_list.loc[already_sold_list["SKU"]
                                                                 == series["SKU"]].iloc[0]["Qty Sold"]
            qspinbox = qtw.QDoubleSpinBox()
            qspinbox.setMaximum(
                series["Weight issued"]-already_returned_weight)
            qspinbox.setValue(
                series["Weight issued"]-already_returned_weight)
            qspinbox.setDecimals(2)
            qspinbox.setToolTip(
                "Previously returned weight + sold weight {} ".format(already_returned_weight))
            qspinbox2 = qtw.QDoubleSpinBox()
            qspinbox2.setMaximum(999999)
            qspinbox2.setMinimum(-999999)
            qspinbox2.setDecimals(2)
            qspinbox2.setToolTip(
                "any unexpected gain/loss in weight during issue period")
            series["Short/Excess Weight"] = qspinbox2
            series["Weight returned"] = qspinbox
            series["Return Date"] = return_date
            return series

        self.issue_approval_list = self.issue_approval_list.apply(
            lambda x: apply_func(x), axis=1)
        # for dicti in self.issue_sheet:
        #     if dicti["APPROVAL NO."] == approval_no:
        #         # check if this approval also exists in issue_return_sheet
        #         already_returned_weight = 0
        #         for return_dicti in self.issue_return_sheet:
        #             if return_dicti["APPROVAL NO."] == approval_no:
        #                 if return_dicti["SKU"] == dicti["SKU"]:
        #                     already_returned_weight = return_dicti["Weight returned"]
        #                     break
        #         qspinbox = qtw.QDoubleSpinBox()
        #         qspinbox.setMaximum(
        #             dicti["Weight issued"]-already_returned_weight)
        #         qspinbox.setValue(
        #             dicti["Weight issued"]-already_returned_weight)
        #         qspinbox.setDecimals(2)
        #         qspinbox.setToolTip(
        #             "Previously returned weight {}".format(already_returned_weight))
        #         temp_item = {"Weight returned": qspinbox,
        #                      "Return Date": return_date}
        #         temp_item.update(deepcopy(dicti))
        #         self.issue_approval_list.append(temp_item)
        self.issue_approval_list = self.issue_approval_list[
            self.schema["issue_return_sheet"]["columns"]]
        read_only = deepcopy(self.schema["issue_return_sheet"]["columns"])
        read_only.remove("Weight returned")
        read_only.remove("Short/Excess Weight")
        self.load_data_in_tabel(self.ui.issue_return_table, self.issue_approval_list, len(
            self.issue_approval_list), self.schema["issue_return_sheet"]["columns"], read_only=read_only, set_cell_widget=["Weight returned", "Short/Excess Weight"])
        self.ui.issue_returned.setEnabled(True)
        self.ui.reset_issue_return.setEnabled(True)

    def add_sku(self):
        # Validate if approval number is unique
        approval_no = self.ui.approval_no_2.value()
        self.ui.approval_no_2.setPalette(self.style().standardPalette())
        self.ui.broker_name.setPalette(self.style().standardPalette())
        self.ui.party_name.setPalette(self.style().standardPalette())
        self.ui.Weight_to_add.setPalette(self.style().standardPalette())
        self.ui.enter_sku_code.setPalette(self.style().standardPalette())
        valid = True
        add = True
        # for row in self.issue_sheet:
        #     if approval_no == row["APPROVAL NO."]:
        #         valid = False
        #         break
        if approval_no in self.issue_sheet["APPROVAL NO."]:
            valid = False
        if valid:
            if approval_no in self.issue_return_sheet["APPROVAL NO."]:
                valid = False
            # for row in self.issue_return_sheet:
            #     if approval_no == row["APPROVAL NO."]:
            #         valid = False
            #         break
        if self.ui.broker_name.text() == '':
            self.ui.broker_name.setPalette(self.red_palette())
            self.ui.statusbar.showMessage("Enter Broker Name")
            add = False
        if self.ui.party_name.text() == '':
            self.ui.party_name.setPalette(self.red_palette())
            self.ui.statusbar.showMessage("Enter Party Name")
            add = False
        if self.ui.Weight_to_add.value() == 0:
            self.ui.Weight_to_add.setPalette(self.red_palette())
            self.ui.statusbar.showMessage("Weight can not be 0")
            add = False
        if not valid:
            self.ui.statusbar.showMessage("Approval No. Invalid")
            self.ui.approval_no_2.setPalette(self.red_palette())
            add = False
        if add:
            # fix broker name and party name
            self.ui.broker_name.setReadOnly(True)
            self.ui.party_name.setReadOnly(True)
            self.ui.dateEdit_2.setReadOnly(True)
            self.ui.approval_no_2.setReadOnly(True)

            item_to_add = self.stock_sheet.iloc[self.sheet_index_add_sku].copy(
            )

            if self.ui.add_sku.text() == "Add":
                self.issue_sheet_skus.append(copy(item_to_add["SKU"]))
                item_to_add.drop(["Weight in Hand",
                                  "Short/Excess Received",
                                  "Issue Qty",
                                  "Sold Qty",
                                  "# Code",
                                  "Remarks",
                                  "Date"], inplace=True)
                item_to_add["Approval Date"] = self.ui.dateEdit_2.text()
                item_to_add["APPROVAL NO."] = self.ui.approval_no_2.value()
                item_to_add["BROKER NAME"] = self.ui.broker_name.text()
                item_to_add["PARTY NAME"] = self.ui.party_name.text()
                item_to_add["Weight issued"] = self.ui.Weight_to_add.value()
                qcombobox = qtw.QComboBox()
                qcombobox.addItems(['No', 'Yes'])
                item_to_add["SELECTION YES/NO"] = qcombobox
                item_dict = {}
                for key in self.schema["issue_sheet"]["columns"]:
                    if item_to_add.get(key):
                        item_dict[key] = item_to_add[key]
                    else:
                        item_dict[key] = nan
                self.ui.issue_items.setEnabled(True)

                self.issue_sheet_items = self.issue_sheet_items.append(
                    item_dict, ignore_index=True)
                affected_row = self.ui.issue_items.rowCount()
                row_count = self.ui.issue_items.rowCount()+1
            else:
                # If update is required
                # affected_row = 0
                row_count = self.ui.issue_items.rowCount()
                self.issue_sheet_items.loc[self.issue_sheet_items["SKU"] ==
                                           item_to_add["SKU"], "Weight issued"] = self.ui.Weight_to_add.value()
                affected_row = self.issue_sheet_items.loc[self.issue_sheet_items["SKU"] == item_to_add["SKU"]].index.tolist()[
                    0]
                # for dicti in self.issue_sheet_items:
                #     if dicti["SKU"] == item_to_add["SKU"]:
                #         dicti["Weight issued"] = self.ui.Weight_to_add.value()
                #         item_dict = deepcopy(dicti)
                #         break
                #     affected_row += 1
                self.ui.add_sku.setText("Add")

            self.load_data_in_tabel(
                self.ui.issue_items, values=self.issue_sheet_items,
                rowcount=row_count,
                columns=self.schema["issue_sheet"]["columns"],
                affected_row=affected_row, read_only=[item for item in self.schema["issue_sheet"]["columns"] if item not in [
                    "SELECTION YES/NO",
                    "SELECTION CRITERIA",
                    "ASKING RATE", "Remarks"]], set_cell_widget=["SELECTION YES/NO"])

            self.ui.approval_issue.setEnabled(True)
            self.ui.Delete_issue_items.setEnabled(True)
            self.ui.reset_issue_items.setEnabled(True)
            self.ui.clear_issue_stock.click()

    def issue_approval(self):
        filepath = qtw.QFileDialog.getSaveFileName(
            self, "Save File", filter=".docx")

        if bool(filepath):
            filepath = filepath[0] + filepath[1]
            self.ui.approval_issue.setEnabled(False)
            self.ui.Delete_issue_items.setEnabled(False)
            self.ui.reset_issue_items.setEnabled(False)

            # load values from table to issue_sheet_items
            for row in range(len(self.issue_sheet_items)):
                self.issue_sheet_items.loc[row, "SELECTION YES/NO"] = self.ui.issue_items.cellWidget(
                    row, 13).currentText()
                try:
                    self.issue_sheet_items.loc[row, "SELECTION CRITERIA"] = self.ui.issue_items.item(
                        row, 14).text()
                except AttributeError:
                    self.issue_sheet_items.loc[row, "SELECTION CRITERIA"] = nan
                try:
                    self.issue_sheet_items.loc[row, "ASKING RATE"] = self.ui.issue_items.item(
                        row, 15).text()
                except AttributeError:
                    self.issue_sheet_items.loc[row, "ASKING RATE"] = nan
                try:
                    self.issue_sheet_items.loc[row, "Remarks"] = self.ui.issue_items.item(
                        row, 18).text()
                except AttributeError:
                    self.issue_sheet_items.loc[row, "Remarks"] = nan
                self.stock_sheet.loc[self.stock_sheet["SKU"] == self.issue_sheet_items.iloc[row]["SKU"], [
                    "Weight in Hand", "Issue Qty"]] += [-self.issue_sheet_items.iloc[row]["Weight issued"], self.issue_sheet_items.iloc[row]["Weight issued"]]
                self.load_data_in_tabel(
                    self.ui.current_stock, values=self.stock_sheet,
                    rowcount=self.ui.current_stock.rowCount(),
                    columns=self.schema["stock_sheet"]["columns"],
                    affected_row=self.stock_sheet.loc[self.stock_sheet["SKU"] == self.issue_sheet_items.iloc[row]["SKU"]].index.tolist()[0])

            self.issue_sheet = self.issue_sheet.append(
                self.issue_sheet_items, ignore_index=True)

            doc = docx.Document("test_project/approval_format.docx")
            p = doc.paragraphs[0]
            p.clear()
            p.add_run('APPROVAL NO.: {} \nTo: {} \nDate: {}'.format(
                self.issue_sheet_items.iloc[0]["APPROVAL NO."], self.issue_sheet_items.iloc[0]["BROKER NAME"], self.issue_sheet_items.iloc[0]["Approval Date"]))

            t = doc.tables[0]
            # t = doc.add_table(len(self.issue_sheet_items)+1, 4)
            # headers = ["Sr. No.", "Particulars", "Weight\n(ct)", "Rate"]
            # for j in range(4):
            #     t.cell(0, j).text = headers[j]
            #     t.cell(0, j).bold = True

            # row_template = deepcopy(t.rows[1])

            for i in range(1, len(self.issue_sheet_items)+1):
                t.add_row()
                for j in range(4):
                    if j == 0:
                        t.cell(i, j).text = str(i)
                    elif j == 1:
                        t.cell(i, j).text = "{} {} {}".format(
                            self.issue_sheet_items.iloc[i-1]["SKU"], self.issue_sheet_items.iloc[i-1]["Shape"], self.issue_sheet_items.iloc[i-1]["Size"])
                    elif j == 2:
                        t.cell(i, j).text = str(
                            self.issue_sheet_items.iloc[i-1]["Weight issued"])
                    elif j == 3:
                        t.cell(i, j).text = str(
                            self.issue_sheet_items.iloc[i-1]["ASKING RATE"])

            try:
                t.style = 'Table-Grid'
            except:
                # May only work in windows
                False

            # print_dlg = qtp.QPrintPreviewDialog()
            # if print_dlg.exec_() == qtw.QDialog.Accepted:

            doc.save(filepath)

            self.ui.broker_name.setReadOnly(False)
            self.ui.party_name.setReadOnly(False)
            self.ui.dateEdit_2.setReadOnly(False)
            self.ui.approval_no_2.setReadOnly(False)
            self.set_approval_no()
            # clear variables
            self.clear_tabel(self.ui.issue_items, disable=True, row_count=0)
            self.sheet_index_add_sku = None
            self.issue_sheet_skus = []
            self.issue_sheet_items = pd.DataFrame(
                columns=self.schema["issue_sheet"]["columns"])

    def delete_item_from_issue_table(self):
        indexes = list(set([item.row()
                       for item in self.ui.issue_items.selectedIndexes()]))
        if len(indexes) == self.ui.issue_items.rowCount():
            self.reset_issue_items()
        else:
            for row in indexes:
                self.issue_sheet_items.drop(
                    self.issue_sheet_items[self.issue_sheet_items["SKU"] == self.ui.issue_items.item(row, 2).text()].index, inplace=True)
                self.issue_sheet_skus.remove(
                    self.ui.issue_items.item(row, 2).text())
            self.load_data_in_tabel(self.ui.issue_items, self.issue_sheet_items, len(self.issue_sheet_items), self.schema["issue_sheet"]["columns"],
                                    read_only=[item for item in self.schema["issue_sheet"]["columns"] if item not in [
                                        "SELECTION YES/NO",
                                        "SELECTION CRITERIA",
                                        "ASKING RATE", "Remarks"]], set_cell_widget=["SELECTION YES/NO"])
        # if len(self.ui.issue_items.selectedIndexes()) == self.ui.issue_items.columnCount() and self.ui.issue_items.selectedIndexes().count(self.ui.issue_items.selectedIndexes()[0]) == len(self.ui.issue_items.selectedIndexes()):
        #     # save item
        #     sheet_index = self.issue_sheet_items[self.ui.issue_items.selectedIndexes()[
        #         0].row()]["stock index"]
        #     self.stock_sheet[sheet_index]["Weight in Hand"] += self.issue_sheet_items[self.ui.issue_items.selectedIndexes()[
        #         0].row()]["Weight issued"]
        #     self.stock_sheet[sheet_index]["Issue Qty"] -= self.issue_sheet_items[self.ui.issue_items.selectedIndexes()[
        #         0].row()]["Weight issued"]
        #     del(self.issue_sheet_items[self.ui.issue_items.selectedIndexes()[
        #         0].row()])
        #     self.ui.issue_items.removeRow(
        #         self.ui.issue_items.selectedIndexes()[0].row())
        # else:
        #     indexes = self.ui.issue_items.selectedIndexes()
        #     for index in indexes:
        #         self.ui.issue_items.setItem(
        #             index.row(), index.column(), qtw.QTableWidgetItem(str('')))

        # self.load_data_in_tabel(self.ui.current_stock, self.issue_sheet, len(
        #     self.issue_sheet), self.schema["stock_sheet"]["columns"])

    def reset_issue_items(self):
        self.issue_sheet_items = pd.DataFrame(
            columns=self.schema["issue_sheet"]["columns"])
        self.issue_sheet_skus = []
        self.ui.broker_name.setReadOnly(False)
        self.ui.party_name.setReadOnly(False)
        self.ui.dateEdit_2.setReadOnly(False)
        self.ui.approval_no_2.setReadOnly(False)
        self.clear_tabel(self.ui.issue_items, disable=True, row_count=0)

        # for index in range(self.ui.issue_items.rowCount()):
        # sheet_index = self.issue_sheet_items[index]["stock index"]
        # self.stock_sheet[sheet_index]["Weight in Hand"] += self.issue_sheet_items[index]["Weight issued"]
        # self.stock_sheet[sheet_index]["Issue Qty"] -= self.issue_sheet_items[index]["Weight issued"]
        # del(self.issue_sheet_items[index])
        # self.ui.issue_items.removeRow(
        # index)

    def find_sku(self):
        self.ui.add_sku.setText("Add")
        self.clear_tabel(self.ui.sku_find, disable=True)
        self.ui.enter_sku_code.setPalette(self.style().standardPalette())
        self.ui.add_sku.setEnabled(False)
        sku_to_find = self.ui.enter_sku_code.text()
        if not bool(sku_to_find):
            return
        sku_found = self.stock_sheet.loc[self.stock_sheet["SKU"]
                                         == sku_to_find]
        self.sheet_index_add_sku = sku_found.index.tolist()[0]
        # count = 0
        # for dicti in self.stock_sheet:
        #     if dicti["SKU"] == sku_to_find:
        #         sku_found = dicti
        #         self.sheet_index_add_sku = count
        #         break
        #     count += 1
        if not sku_found.empty:
            self.load_data_in_tabel(self.ui.sku_find, values=sku_found,
                                    rowcount=1, columns=self.schema["stock_sheet"]["columns"])
            sku_found = sku_found.iloc[0]
            self.ui.sku_find.setEnabled(True)
            self.ui.add_sku.setEnabled(True)
            self.ui.Weight_to_add.setPalette(self.style().standardPalette())
            self.ui.Weight_to_add.setEnabled(True)
            self.ui.Weight_to_add.setMaximum(
                sku_found["Weight in Hand"])
            self.ui.Weight_to_add.setValue(sku_found["Weight in Hand"])

            # If item is already in to be issue items sheet change add to update
            if sku_found["SKU"] in self.issue_sheet_skus:
                self.ui.add_sku.setText("Update")
        else:
            self.ui.enter_sku_code.setPalette(self.red_palette())

    def add_stock(self):
        entry_dict = {}
        # cell_widgets = ["Date", "Item", "Mine", "Shape", "Original weight", "Short/Excess Received"]
        for j in range(self.ui.current_entry_add_stock.columnCount()):
            column = self.ui.current_entry_add_stock.horizontalHeaderItem(
                j).text()
            try:
                entry_dict[column] = self.ui.current_entry_add_stock.item(
                    0, j).text()
            except AttributeError:
                entry_dict[column] = self.ui.current_entry_add_stock.item(0, j)
        entry_dict["Date"] = self.ui.current_entry_add_stock.cellWidget(
            0, 1).text()
        entry_dict["Item"] = self.ui.current_entry_add_stock.cellWidget(
            0, 3).currentText()
        entry_dict["Mine"] = self.ui.current_entry_add_stock.cellWidget(
            0, 6).currentText()
        entry_dict["Shape"] = self.ui.current_entry_add_stock.cellWidget(
            0, 4).currentText()
        entry_dict["Original weight"] = self.ui.current_entry_add_stock.cellWidget(
            0, 10).value()
        entry_dict["Short/Excess Received"] = self.ui.current_entry_add_stock.cellWidget(
            0, 12).value()
        entry_dict["Weight in Hand"] = entry_dict["Original weight"] + \
            entry_dict["Short/Excess Received"]
        entry_dict["Issue Qty"] = 0
        entry_dict["Sold Qty"] = 0
        self.stock_sheet = self.stock_sheet.append(
            entry_dict, ignore_index=True)
        # self.stock_sheet.sort_values("Approval Date", inplace=True, ignore_index=True)
        # self.ui.current_entry_add_stock.clear()
        self.load_data_in_tabel(self.ui.current_stock, self.stock_sheet, len(
            self.stock_sheet), list(entry_dict), len(self.stock_sheet)-1)
        self.init_add_stock()
        # tabel_columns = []
        # for i in range(self.ui.current_entry_add_stock.columnCount()):
        #     tabel_columns.append(
        #         self.ui.current_entry_add_stock.horizontalHeaderItem(i).text())
        # self.ui.current_entry_add_stock.clear()
        # self.ui.current_entry_add_stock.setHorizontalHeaderLabels(
        #     tabel_columns)

    def init_add_stock(self):
        # prepares the table to accept new data
        add_stock_format = pd.DataFrame(
            columns=self.schema["stock_sheet"]["columns"])
        table_date = qtw.QDateEdit(date=qtc.QDate(
            date.today().year, date.today().month, date.today().day))
        table_date.setCalendarPopup(True)

        table_item = qtw.QComboBox()
        table_item.addItems(list(self.schema["items"]))
        table_item.setEditable(True)
        table_item.completer().setCompletionMode(qtw.QCompleter.PopupCompletion)
        table_item.setInsertPolicy(qtw.QComboBox.NoInsert)
        table_item.currentTextChanged.connect(
            self.handler_add_stock_sku_from_combo_box)

        table_mine = qtw.QComboBox()
        table_mine.addItems(list(self.schema["mines"]))
        table_mine.setEditable(True)
        table_mine.completer().setCompletionMode(qtw.QCompleter.PopupCompletion)
        table_mine.setInsertPolicy(qtw.QComboBox.NoInsert)
        table_mine.currentTextChanged.connect(
            self.handler_add_stock_sku_from_combo_box)

        table_shape = qtw.QComboBox()
        table_shape.addItems(self.schema["shapes"])
        table_shape.setEditable(True)
        table_shape.completer().setCompletionMode(qtw.QCompleter.PopupCompletion)
        table_shape.setInsertPolicy(qtw.QComboBox.NoInsert)

        table_weight = qtw.QDoubleSpinBox()
        table_weight.setMaximum(99999999)
        table_weight.setDecimals(2)
        table_weight.valueChanged.connect(
            self.handler_add_stock_original_weight)

        table_short_excess = qtw.QDoubleSpinBox()
        table_short_excess.setMaximum(0)
        table_short_excess.setMinimum(0)
        table_short_excess.setDecimals(2)
        table_short_excess.valueChanged.connect(
            self.handler_add_stock_short_excess_received)

        # set SKU
        sku = self.schema["items"][table_item.currentText(
        )]+self.schema["mines"][table_mine.currentText()]
        max_num = self.stock_sheet.loc[self.stock_sheet["SKU"].astype(
            str).str[:5] == sku]["SKU"].astype(str).str[5:].astype(int).max()
        max_num = 0 if pd.isna(max_num) else max_num
        sku = sku+str(max_num+1)

        add_stock_format = add_stock_format.append({
            "SKU": sku,
            "Date": table_date,
            "Item": table_item,
            "Mine": table_mine,
            "Shape": table_shape,
            "Original weight": table_weight,
            "Short/Excess Received": table_short_excess,
            "Issue Qty": 0,
            "Sold Qty": 0
        }, ignore_index=True)
        self.load_data_in_tabel(
            self.ui.current_entry_add_stock, add_stock_format, 1, self.schema["stock_sheet"]["columns"], read_only=["SKU", "Weight in Hand", "Issue Qty", "Sold Qty"],
            set_cell_widget=["Date", "Item", "Mine", "Shape", "Original weight", "Short/Excess Received"])

    def confirm_edit(self, tab="stock"):
        if tab == "stock":
            for i in range(self.ui.find_stock_preview.rowCount()):
                temp_series = pd.Series(
                    index=self.schema["stock_sheet"]["columns"])
                for j in range(self.ui.find_stock_preview.columnCount()):
                    temp_series[self.ui.find_stock_preview.horizontalHeaderItem(
                        j).text()] = self.ui.find_stock_preview.item(i, j).text()

                self.stock_sheet.loc[int(self.ui.find_stock_preview.verticalHeaderItem(
                    i).text())] = temp_series.copy()
            self.load_data_in_tabel(self.ui.current_stock, self.stock_sheet, len(
                self.stock_sheet), self.schema["stock_sheet"]["columns"])

        elif tab == "issue":
            for i in range(self.ui.issue_stock_preview.rowCount()):
                temp_series = pd.Series(
                    index=self.schema["issue_sheet"]["columns"])
                for j in range(1, self.ui.issue_stock_preview.columnCount()):
                    temp_series[self.ui.issue_stock_preview.horizontalHeaderItem(
                        j).text()] = self.ui.issue_stock_preview.item(i, j).text()
                self.issue_sheet.loc[int(self.ui.issue_stock_preview.verticalHeaderItem(
                    i).text())] = temp_series.copy()

                ##### Add recalculate stock_sheet here ####

    def reset_selected(self, event):
        self.ui.current_stock_find_entry.clearSelection()
        self.ui.find_stock_preview.clearSelection()
        self.ui.current_entry_add_stock.clearSelection()
        self.ui.current_stock.clearSelection()

    def reset_selected_issue(self, event):
        self.ui.issue_stock_preview.clearSelection()
        self.ui.issued_stock_find_entry.clearSelection()
        self.ui.sku_find.clearSelection()
        self.ui.issue_items.clearSelection()
        if not self.ui.sku_find.isEnabled():
            # self.ui.enter_sku_code.clear()
            self.ui.Weight_to_add.clear()
            self.ui.Weight_to_add.setMaximum(0)
        if not bool(self.ui.enter_sku_code.text()):
            # self.ui.approval_no_2.palette()
            self.ui.enter_sku_code.setPalette(self.style().standardPalette())

    def issue_find_enter_event(self, event):
        # Enter key pressed
        if event.key() == qtc.Qt.Key_Return or event.key() == qtc.Qt.Key_Enter:
            if len(self.ui.issued_stock_find_entry.selectedIndexes()) == 1:
                item = self.ui.issued_stock_find_entry.itemFromIndex(
                    self.ui.issued_stock_find_entry.selectedIndexes()[0])
                if item is None:
                    self.ui.issued_stock_find_entry.edit(
                        self.ui.issued_stock_find_entry.selectedIndexes()[0])
                else:
                    self.find_item_in_stock_sheet(sheet="issue_sheet")
        qtw.QTableWidget.keyPressEvent(self.ui.issued_stock_find_entry, event)

    def issue_preview_enter_event(self, event):
        # Enter key pressed
        if event.key() == qtc.Qt.Key_Return or event.key() == qtc.Qt.Key_Enter:
            index = self.ui.issue_stock_preview.selectedIndexes()
            if len(index) == 1:
                index = index[0]
                item = self.ui.issue_stock_preview.itemFromIndex(index)

                self.stock_sheet[int(self.ui.issue_stock_preview.item(index.row(
                ), 0).text())-1][self.ui.issue_stock_preview.horizontalHeaderItem(index.column()).text()] = item.text()
                self.load_data_in_tabel(self.ui.current_stock, self.stock_sheet, len(
                    self.stock_sheet), self.schema["stock_sheet"]["columns"])
        if event.key() == qtc.Qt.Key_Delete:
            self.delete_entry_from_tabel(tab="issue")

        qtw.QTableWidget.keyPressEvent(self.ui.issue_stock_preview, event)

    def find_enter_event(self, event):
        # Enter key pressed
        if event.key() == qtc.Qt.Key_Return or event.key() == qtc.Qt.Key_Enter:
            if len(self.ui.current_stock_find_entry.selectedIndexes()) == 1:
                item = self.ui.current_stock_find_entry.itemFromIndex(
                    self.ui.current_stock_find_entry.selectedIndexes()[0])
                if item is None:
                    self.ui.current_stock_find_entry.edit(
                        self.ui.current_stock_find_entry.selectedIndexes()[0])
                else:
                    self.find_item_in_stock_sheet(sheet="stock_sheet")
        qtw.QTableWidget.keyPressEvent(self.ui.current_stock_find_entry, event)

    def preview_enter_event(self, event):
        # Enter key pressed
        if event.key() == qtc.Qt.Key_Return or event.key() == qtc.Qt.Key_Enter:
            index = self.ui.find_stock_preview.selectedIndexes()
            if len(index) == 1:
                index = index[0]
                item = self.ui.find_stock_preview.itemFromIndex(index)
                self.stock_sheet.loc[int(self.ui.find_stock_preview.verticalHeaderItem(index.row()).text(
                )), self.ui.find_stock_preview.horizontalHeaderItem(index.column()).text()] = item.text()
                # self.stock_sheet[int(self.ui.find_stock_preview.item(index.row(
                # ), 0).text())-1][self.ui.find_stock_preview.horizontalHeaderItem(index.column()).text()] = item.text()
                self.load_data_in_tabel(self.ui.current_stock, self.stock_sheet, len(
                    self.stock_sheet), self.schema["stock_sheet"]["columns"])
        if event.key() == qtc.Qt.Key_Delete:
            self.delete_entry_from_tabel(tab="stock")

        qtw.QTableWidget.keyPressEvent(self.ui.find_stock_preview, event)

    def reset_preview(self, tab="stock"):
        if tab == "stock":
            # for dicti in self.original:

            #     index = copy(int(dicti["sheet_index"]))
            #     temp_dict = deepcopy(dicti)
            #     del(temp_dict["sheet_index"])
            #     self.stock_sheet[index-1] = deepcopy(temp_dict)
            self.stock_sheet = self.stock_sheet.append(self.original)
            self.load_data_in_tabel(self.ui.current_stock, self.stock_sheet, len(
                self.stock_sheet), self.schema["stock_sheet"]["columns"])
            self.load_data_in_tabel(self.ui.find_stock_preview, self.original, len(
                self.original), self.schema["stock_sheet"]["columns"])

        elif tab == "issue":
            # for dicti in self.original_issue:

            #     index = copy(int(dicti["sheet_index"]))
            #     temp_dict = deepcopy(dicti)
            #     del(temp_dict["sheet_index"])
            #     self.issue_sheet[index-1] = deepcopy(temp_dict)
            self.issue_sheet = self.issue_sheet.append(self.original_issue)
            self.load_data_in_tabel(self.ui.issue_stock_preview, self.original_issue, len(
                self.original_issue), self.schema["issue_sheet"]["columns"])

    def clear_tabel(self, tabel, disable=False, row_count=1):

        assert(isinstance(tabel, qtw.QTableWidget))
        tabel_columns = []
        for i in range(tabel.columnCount()):
            tabel_columns.append(tabel.horizontalHeaderItem(i).text())
        tabel.clear()
        tabel.setHorizontalHeaderLabels(tabel_columns)
        tabel.setRowCount(row_count)
        if disable == True:
            tabel.setEnabled(False)

    def stock_tab_button_set_False(self):
        self.ui.clear_button.setEnabled(False)
        self.ui.delete_entry.setEnabled(False)
        self.ui.reset_find_preview.setEnabled(False)
        self.ui.confirm_edit.setEnabled(False)

    def issue_tab_button_set_false(self):
        self.ui.clear_button_2.setEnabled(False)
        self.ui.delete_entry_2.setEnabled(False)
        self.ui.edit_button_2.setEnabled(False)
        self.ui.Reset_issued_preview.setEnabled(False)

    def delete_entry_from_tabel(self, tab="stock"):
        if tab == "stock":
            # if len(self.ui.find_stock_preview.selectedIndexes()) == self.ui.find_stock_preview.columnCount() and self.ui.find_stock_preview.selectedIndexes().count(self.ui.find_stock_preview.selectedIndexes()[0]) == len(self.ui.find_stock_preview.selectedIndexes()):
            #     # deleting an entire row
            #     temp_dict = {}
            #     for i in range(self.ui.find_stock_preview.columnCount()):
            #         temp_dict[self.ui.find_stock_preview.horizontalHeaderItem(i).text()] = self.ui.find_stock_preview.item(
            #             self.ui.find_stock_preview.selectedIndexes()[0].row(), i).text()
            #     self.deleted.append(self.stock_sheet.pop(
            #         int(temp_dict["sheet_index"])-1))
            #     self.load_data_in_tabel(self.ui.current_stock, self.stock_sheet, len(
            #         self.stock_sheet), self.schema["stock_sheet"]["columns"])
            #     self.ui.find_stock_preview.removeRow(
            #         self.ui.find_stock_preview.selectedIndexes()[0].row())
            # else:
            indexes = self.ui.find_stock_preview.selectedIndexes()
            for index in indexes:
                self.ui.find_stock_preview.setItem(
                    index.row(), index.column(), qtw.QTableWidgetItem(str('')))
                self.stock_sheet.loc[int(self.ui.find_stock_preview.verticalHeaderItem(index.row()).text(
                )), self.ui.find_stock_preview.horizontalHeaderItem(index.column()).text()] = nan
                # self.stock_sheet[int(self.ui.find_stock_preview.item(index.row(
                # ), 0).text())-1][self.ui.find_stock_preview.horizontalHeaderItem(index.column()).text()] = ''
            # self.stock_sheet.dropna(how='all', axis=0, inplace=True)
            # self.stock_sheet.sort_values("Approval Date", inplace=True, ignore_index=True)

            self.load_data_in_tabel(self.ui.current_stock, self.stock_sheet, len(
                self.stock_sheet), self.schema["stock_sheet"]["columns"])

        elif tab == "issue":
            # if len(self.ui.issue_stock_preview.selectedIndexes()) == self.ui.issue_stock_preview.columnCount() and self.ui.issue_stock_preview.selectedIndexes().count(self.ui.issue_stock_preview.selectedIndexes()[0]) == len(self.ui.issue_stock_preview.selectedIndexes()):
            #     # save item
            #     temp_dict = {}
            #     for i in range(self.ui.issue_stock_preview.columnCount()):
            #         temp_dict[self.ui.issue_stock_preview.horizontalHeaderItem(i).text()] = self.ui.issue_stock_preview.item(
            #             self.ui.issue_stock_preview.selectedIndexes()[0].row(), i).text()
            #     self.deleted.append(self.issue_sheet.pop(
            #         int(temp_dict["sheet_index"])-1))
            #     # self.load_data_in_tabel(self.ui.current_stock, self.stock_sheet, len(
            #     #     self.stock_sheet), self.schema["stock_sheet"]["columns"])
            #     self.ui.issue_stock_preview.removeRow(
            #         self.ui.issue_stock_preview.selectedIndexes()[0].row())
            # else:
            indexes = self.ui.issue_stock_preview.selectedIndexes()
            for index in indexes:
                self.ui.issue_stock_preview.setItem(
                    index.row(), index.column(), qtw.QTableWidgetItem(str('')))
                self.issue_sheet.loc[int(self.ui.issue_stock_preview.verticalHeaderItem(index.row()).text(
                )), self.ui.issue_stock_preview.horizontalHeaderItem(index.column()).text()] = nan
                # self.issue_sheet[int(self.ui.issue_stock_preview.item(index.row(
                # ), 0).text())-1][self.ui.issue_stock_preview.horizontalHeaderItem(index.column()).text()] = ''

    def find_item_in_stock_sheet(self, sheet):

        if sheet == "stock_sheet":
            self.clear_tabel(self.ui.find_stock_preview)
            self.original = []
            self.ui.current_stock_find_entry.clearSelection()
            search_dict = {}
            for i in range(len(self.schema["stock_sheet"]["columns"])):
                try:
                    if self.ui.current_stock_find_entry.item(0, i).text() != "":
                        search_dict[self.schema["stock_sheet"]["columns"][i]
                                    ] = self.ui.current_stock_find_entry.item(0, i).text()
                except AttributeError:
                    False
            if not bool(search_dict):
                return

            self.original = self.stock_sheet.loc[(
                self.stock_sheet[list(search_dict)] == pd.Series(search_dict)).all(axis=1)]
            if not self.original.empty:
                self.ui.find_stock_preview.setEnabled(True)
                self.load_data_in_tabel(self.ui.find_stock_preview, self.original, len(
                    self.original), self.schema["stock_sheet"]["columns"])
                self.ui.clear_button.setEnabled(True)
                self.ui.reset_find_preview.setEnabled(True)
                self.ui.confirm_edit.setEnabled(True)
                # self.ui.edit_button.setEnabled(True)
                self.ui.delete_entry.setEnabled(True)

        elif sheet == "issue_sheet":
            self.clear_tabel(self.ui.issue_stock_preview)
            self.original_issue = []
            self.ui.issued_stock_find_entry.clearSelection()
            search_dict = {}
            for i in range(len(self.schema["issue_sheet"]["columns"])):
                try:
                    if self.ui.issued_stock_find_entry.item(0, i).text() != "":
                        search_dict[self.schema["issue_sheet"]["columns"][i]
                                    ] = self.ui.issued_stock_find_entry.item(0, i).text()
                except AttributeError:
                    False
            if not bool(search_dict):
                return
            self.original_issue = self.issue_sheet.loc[(
                self.issue_sheet[list(search_dict)] == pd.Series(search_dict)).all(axis=1)]
            if not self.original_issue.empty:
                self.ui.issue_stock_preview.setEnabled(True)
                self.load_data_in_tabel(self.ui.issue_stock_preview, self.original_issue, len(
                    self.original_issue), self.schema["issue_sheet"]["columns"])
                self.ui.clear_button_2.setEnabled(True)
                self.ui.Reset_issued_preview.setEnabled(True)
                self.ui.edit_button_2.setEnabled(True)
                self.ui.delete_entry_2.setEnabled(True)

    def load_data_in_tabel(self, tabel, values, rowcount, columns, affected_row=None, read_only=[], set_cell_widget=[]):
        # affected_row =
        assert(isinstance(tabel, qtw.QTableWidget))
        assert isinstance(values, pd.DataFrame)
        tabel.setRowCount(rowcount)
        tabel.setColumnCount(
            len(columns))
        tabel.setHorizontalHeaderLabels(
            columns)
        tabel.setVerticalHeaderLabels(values.index.astype(str).tolist())
        if affected_row == None:
            # load entire data in table
            if bool(read_only):
                for row, series in values.reset_index(drop=True).iterrows():
                    for col, key in enumerate(columns):
                        # here values is one dict
                        item = series[key] if key in set_cell_widget else qtw.QTableWidgetItem(
                            str(series[key]) if pd.notna(series[key]) else '')
                        if key in read_only:
                            item.setFlags(qtc.Qt.ItemIsSelectable | qtc.Qt.ItemIsDragEnabled |
                                          qtc.Qt.ItemIsDropEnabled | qtc.Qt.ItemIsUserCheckable | qtc.Qt.ItemIsEnabled)
                        if key in set_cell_widget:
                            tabel.setCellWidget(row, col, item)
                        else:
                            tabel.setItem(row, col,
                                          item)
            else:
                for row, series in values.reset_index(drop=True).iterrows():
                    for col, key in enumerate(columns):
                        tabel.setItem(row, col,
                                      qtw.QTableWidgetItem(str(series[key]) if pd.notna(series[key]) else ''))

        else:
            # load only the respective data
            values = values.iloc[affected_row]
            if bool(read_only):
                for col, key in enumerate(columns):
                    # here values is one dict
                    item = values[key] if key in set_cell_widget else qtw.QTableWidgetItem(
                        str(values[key]) if pd.notna(values[key]) else '')
                    # item = qtw.QTableWidgetItem('' if pd.isna(
                    #     values[key]) else str(values[key]))
                    # if key in read_only:
                    #     item.setFlags(qtc.Qt.ItemIsSelectable | qtc.Qt.ItemIsDragEnabled |
                    #                   qtc.Qt.ItemIsDropEnabled | qtc.Qt.ItemIsUserCheckable | qtc.Qt.ItemIsEnabled)
                    # tabel.setItem(affected_row, col,
                    #               item)
                    if key in read_only:
                        item.setFlags(qtc.Qt.ItemIsSelectable | qtc.Qt.ItemIsDragEnabled |
                                      qtc.Qt.ItemIsDropEnabled | qtc.Qt.ItemIsUserCheckable | qtc.Qt.ItemIsEnabled)
                    if key in set_cell_widget:
                        tabel.setCellWidget(affected_row, col,
                                            item)
                    else:
                        tabel.setItem(affected_row, col,
                                      item)
            else:
                for col, key in enumerate(columns):
                    # here values is one dict
                    item = values[key] if key in set_cell_widget else qtw.QTableWidgetItem(
                        str(values[key]) if pd.notna(values[key]) else '')
                    if key in set_cell_widget:
                        tabel.setCellWidget(row, col, item)
                    else:
                        tabel.setItem(affected_row, col,
                                      item)

        tabel.horizontalHeader(
        ).resizeSections(qtw.QHeaderView.ResizeToContents)

    def save_project(self, save_path):
        self.stock_sheet.to_excel(os.path.join(
            save_path, self.schema["stock_sheet"]["filename"]), index=False)
        self.issue_sheet.to_excel(os.path.join(
            save_path, self.schema["issue_sheet"]["filename"]), index=False)
        self.issue_return_sheet.to_excel(os.path.join(
            save_path, self.schema["issue_return_sheet"]["filename"]), index=False)
        self.sale_sheet.to_excel(os.path.join(
            save_path, self.schema["sale_sheet"]["filename"]), index=False)

    def open_project(self, path=False):

        try:
            if path == False:
                self.project_path = qtw.QFileDialog.getExistingDirectory(
                    self, "Select Directory")
            if bool(self.project_path):
                f = open(os.path.join(self.project_path, "schema.json"))
                self.schema = json.loads(f.read())
                f.close()

                self.stock_sheet = pd.read_excel(os.path.join(
                    self.project_path, self.schema["stock_sheet"]["filename"]))
                self.stock_sheet = self.stock_sheet[self.schema["stock_sheet"]["columns"]]
                self.issue_sheet = pd.read_excel(os.path.join(
                    self.project_path, self.schema["issue_sheet"]["filename"]))
                self.issue_sheet = self.issue_sheet[self.schema["issue_sheet"]["columns"]]
                self.issue_return_sheet = pd.read_excel(os.path.join(
                    self.project_path, self.schema["issue_return_sheet"]["filename"]))
                self.issue_return_sheet = self.issue_return_sheet[
                    self.schema["issue_return_sheet"]["columns"]]
                self.sale_sheet = pd.read_excel(os.path.join(
                    self.project_path, self.schema["sale_sheet"]["filename"]))
                self.sale_sheet = self.sale_sheet[self.schema["sale_sheet"]["columns"]]

                # stock tab
                self.load_data_in_tabel(self.ui.current_stock, self.stock_sheet, len(
                    self.stock_sheet), self.schema["stock_sheet"]["columns"])
                self.load_data_in_tabel(self.ui.current_stock_find_entry, pd.DataFrame(
                    columns=self.schema["stock_sheet"]["columns"]), 1, self.schema["stock_sheet"]["columns"])
                self.load_data_in_tabel(
                    self.ui.find_stock_preview, pd.DataFrame(columns=self.schema["stock_sheet"]["columns"]), 1, self.schema["stock_sheet"]["columns"])
                self.init_add_stock()

                # issue tab
                self.load_data_in_tabel(self.ui.issued_stock_find_entry, pd.DataFrame(
                    columns=self.schema["issue_sheet"]["columns"]), 1, self.schema["issue_sheet"]["columns"])
                self.load_data_in_tabel(
                    self.ui.issue_stock_preview, pd.DataFrame(columns=self.schema["issue_sheet"]["columns"]), 1, self.schema["issue_sheet"]["columns"])
                self.load_data_in_tabel(self.ui.sku_find, pd.DataFrame(columns=self.schema["issue_sheet"]["columns"]), 1,
                                        self.schema["stock_sheet"]["columns"])
                self.load_data_in_tabel(self.ui.issue_items, pd.DataFrame(columns=self.schema["issue_sheet"]["columns"]),
                                        0, self.schema["issue_sheet"]["columns"])

                # variabels
                self.issue_sheet_items = pd.DataFrame(
                    columns=self.schema["issue_sheet"]["columns"])

                # qtw.QHeaderView.setSectionResizeMode(qtw.QHeaderView.Stretch)

                self.ui.tabWidget.setTabEnabled(3, True)
                self.ui.tabWidget.setTabEnabled(2, True)
                self.ui.tabWidget.setTabEnabled(1, True)
                self.ui.tabWidget.setTabEnabled(0, True)
                self.ui.actionSave_As.setEnabled(True)
                self.ui.actionSave.setEnabled(True)

                self.ui.issued_stock_find_entry.horizontalHeader(
                ).resizeSections(qtw.QHeaderView.ResizeToContents)
                self.ui.issue_stock_preview.horizontalHeader(
                ).resizeSections(qtw.QHeaderView.ResizeToContents)
                self.ui.sku_find.horizontalHeader().resizeSections(
                    qtw.QHeaderView.ResizeToContents)
                self.ui.issue_items.horizontalHeader().resizeSections(
                    qtw.QHeaderView.ResizeToContents)

                self.ui.current_stock_find_entry.horizontalHeader(
                ).resizeSections(qtw.QHeaderView.ResizeToContents)
                self.ui.find_stock_preview.horizontalHeader().resizeSections(
                    qtw.QHeaderView.ResizeToContents)
                self.ui.current_entry_add_stock.horizontalHeader(
                ).resizeSections(qtw.QHeaderView.ResizeToContents)
                self.ui.current_stock.horizontalHeader().resizeSections(
                    qtw.QHeaderView.ResizeToContents)

                # Set Approval No.
                self.set_approval_no()
            self.ui.actionOpen_Project.setEnabled(False)
        except:
            msg = qtw.QMessageBox()
            msg.setIcon(qtw.QMessageBox.Critical)
            msg.setText("Error")
            msg.setInformativeText('Not a valid project folder')
            msg.setWindowTitle("Error")
            msg.exec_()

    def handler_add_stock_original_weight(self, value):
        w = self.ui.current_entry_add_stock.cellWidget(0, 12)
        w.setMaximum(999999)
        w.setMinimum(-999999)
        w.setDecimals(2)
        self.ui.current_entry_add_stock.item(0, 11).setText(str(value))

    def handler_add_stock_short_excess_received(self, value):
        original_weight = self.ui.current_entry_add_stock.cellWidget(
            0, 10).value()
        self.ui.current_entry_add_stock.item(
            0, 11).setText(str(original_weight+value))

    def handler_add_stock_sku_from_combo_box(self, value):
        # set SKU
        try:
            sku = self.schema["items"][self.ui.current_entry_add_stock.cellWidget(0, 3).currentText(
            )]+self.schema["mines"][self.ui.current_entry_add_stock.cellWidget(0, 6).currentText()]
            max_num = self.stock_sheet.loc[self.stock_sheet["SKU"].astype(
                str).str[:5] == sku]["SKU"].astype(str).str[5:].astype(int).max()
            max_num = 0 if pd.isna(max_num) else max_num
            sku = sku+str(max_num+1)
            self.ui.current_entry_add_stock.item(0, 0).setText(sku)
        except KeyError:
            # intermediate typing error
            False

    def set_approval_no(self):
        m1 = self.issue_sheet["APPROVAL NO."].max()
        m1 = m1 if pd.notna(m1) else 0
        m2 = self.issue_return_sheet["APPROVAL NO."].max()
        m2 = m2 if pd.notna(m2) else 0
        m1 = max(m1, m2)
        if self.schema["approval_no"] <= m1:
            self.schema["approval_no"] = m1+1
        self.ui.approval_no_2.setValue(self.schema["approval_no"])
        self.schema["approval_no"] += 1

    def New_project(self):
        self.project_path = qtw.QFileDialog.getExistingDirectory(
            self, "Select Directory")
        if bool(self.project_path):
            f = open("schema.json")
            self.schema = json.loads(f.read())
            f.close()
            f = open(os.path.join(self.project_path, "schema.json"), "w")
            f.write(json.dumps(self.schema))
            f.close()
            pd.DataFrame(columns=self.schema["stock_sheet"]["columns"]).to_excel(os.path.join(
                self.project_path, self.schema["stock_sheet"]["filename"]), index=False)
            pd.DataFrame(columns=self.schema["issue_sheet"]["columns"]).to_excel(os.path.join(
                self.project_path, self.schema["issue_sheet"]["filename"]), index=False)
            pd.DataFrame(columns=self.schema["issue_return_sheet"]["columns"]).to_excel(os.path.join(
                self.project_path, self.schema["issue_return_sheet"]["filename"]), index=False)
            pd.DataFrame(columns=self.schema["sale_sheet"]["columns"]).to_excel(os.path.join(
                self.project_path, self.schema["sale_sheet"]["filename"]), index=False)

            self.open_project(path=True)
        # self.ui.tabWidget.setTabEnabled(3, True)
        # self.ui.tabWidget.setTabEnabled(2, True)
        # self.ui.tabWidget.setTabEnabled(1, True)
        # self.ui.tabWidget.setTabEnabled(0, True)
        # self.ui.actionSave_As.setEnabled(True)
        # self.ui.actionSave.setEnabled(True)

    def exit(self):
        if self.saved == False:
            msg = qtw.QMessageBox()
            msg.setText("Please save before closing")
            msg.setStandardButtons(qtw.QMessageBox.Save |
                                   qtw.QMessageBox.Discard)
            msg.setModal(True)
            msg.buttonClicked.connect(self.save_warning)
            msg.exec_()

            # print(retval)
        else:
            os._exit(0)

    def save_warning(self, item):
        if item.text() == "Save":
            self.save_project(self.project_path)


if __name__ == "__main__":
    try:
        app = qtw.QApplication([])
        window = MainWindow()
        window.showMaximized()
        app.exec()
    except Exception as e:
        msg = qtw.QMessageBox()
        msg.setIcon(qtw.QMessageBox.Critical)
        msg.setText("Critical Error Occured")
        msg.setInformativeText(str(e))
        msg.setWindowTitle("Error")
        msg.exec_()
